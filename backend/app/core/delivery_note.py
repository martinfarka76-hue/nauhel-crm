"""
Automatické generování Dodacího listu (Word .docx) ze šablony s tokeny
(app/templates/dodaci_list_template.docx) - vzniká při přechodu Dealu
na stav "Vyrobeno". Šablona obsahuje pole k ručnímu vyplnění při
fyzickém předání (stav dodávky, podpisy) - ty se nepředvyplňují,
kromě checkboxu "Kompletní dodávka", který je v šabloně už zaškrtnutý.
"""
import io
import logging
from datetime import date
from pathlib import Path

from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.models.deal import Deal
from app.models.company import Company
from app.models.contact import Contact
from app.models.calculation import Calculation
from app.models.calculation_item import CalculationItem
from app.models.user import User
from app.models.enums import ItemCategory

logger = logging.getLogger("nauhel_crm.delivery_note")

TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "dodaci_list_template.docx"


def _replace_placeholder_in_cell(cell, placeholder: str, value: str) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value or "")
    # Buňka může obsahovat vnořenou tabulku (např. hlavička dokumentu) -
    # python-docx neprochází vnořené tabulky automaticky, musíme rekurzivně.
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                _replace_placeholder_in_cell(nested_cell, placeholder, value)


def _replace_placeholder(doc: DocxDocument, placeholder: str, value: str) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_placeholder_in_cell(cell, placeholder, value)


def _find_specification_table(doc: DocxDocument):
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = " ".join(c.text for c in table.rows[0].cells)
            if "Popis" in header_text and "Množství" in header_text:
                return table
    return None


def generate_delivery_note(
    db: Session, deal: Deal, company: Company | None, contact: Contact | None, owner: User | None
) -> bytes:
    """Vygeneruje Word dokument dodacího listu a vrátí ho jako bajty."""
    doc = DocxDocument(str(TEMPLATE_PATH))

    cislo = ""
    if deal.sharepoint_folder_year and deal.sharepoint_folder_number:
        cislo = f"{deal.sharepoint_folder_year}_{deal.sharepoint_folder_number:03d}"

    ico_dic = ""
    if company:
        ico_dic = f"{company.ico or ''} / {company.dic or ''}".strip(" /")

    replacements = {
        "{{CISLO}}": cislo,
        "{{DATUM_EXPEDICE}}": date.today().strftime("%d.%m.%Y"),
        "{{OBJEDNAVKA_CISLO}}": cislo,
        "{{VYSTAVIL}}": owner.full_name if owner else "",
        "{{PRIJEMCE_FIRMA}}": company.name if company else "",
        "{{PRIJEMCE_ICO_DIC}}": ico_dic,
        "{{PRIJEMCE_ADRESA}}": company.address if company else "",
        "{{PRIJEMCE_KONTAKT}}": f"{contact.first_name} {contact.last_name}" if contact else "",
        "{{DODACI_FIRMA}}": company.name if company else "",
        "{{DODACI_ULICE}}": company.address if company else "",
        "{{DODACI_MESTO}}": "",
        "{{DODACI_TELEFON}}": contact.phone if contact else "",
    }
    for placeholder, value in replacements.items():
        _replace_placeholder(doc, placeholder, str(value))

    # Aktivní kalkulace + její materiálové položky
    calc = db.query(Calculation).filter(Calculation.deal_id == deal.id, Calculation.is_active.is_(True)).first()
    material_items = []
    if calc:
        items = (
            db.query(CalculationItem)
            .filter(CalculationItem.calculation_id == calc.id)
            .order_by(CalculationItem.display_order)
            .all()
        )
        material_items = [i for i in items if i.category == ItemCategory.MATERIAL]

    spec_table = _find_specification_table(doc)
    if spec_table and material_items:
        # poslední řádek je "CELKEM" - datové řádky jsou mezi hlavičkou a ním
        data_rows = spec_table.rows[1:-1]

        while len(data_rows) < len(material_items):
            spec_table.add_row()
            data_rows = spec_table.rows[1:-1]

        for idx, item in enumerate(material_items):
            row = data_rows[idx]
            row.cells[0].text = str(idx + 1)
            row.cells[1].text = item.name
            row.cells[2].text = ""  # Délky - k ručnímu doplnění
            row.cells[3].text = str(item.quantity)
            row.cells[4].text = item.unit or ""
            row.cells[5].text = ""

        for idx in range(len(material_items), len(data_rows)):
            for cell in data_rows[idx].cells:
                cell.text = ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
